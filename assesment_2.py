import streamlit as st
import pandas as pd
import numpy as np
import re
from langchain_huggingface import HuggingFaceEndpoint
from langchain import PromptTemplate, LLMChain
from docx import Document
from dotenv import load_dotenv
import os
import io 

#loading HF API key from environment variables
load_dotenv()
HF_API_KEY = os.getenv("HF_API_KEY")

# Load environment variablesload_dotenv()
HF_API_KEY = os.getenv("HF_API_KEY")    
if not HF_API_KEY:
    st.error("HF_API_KEY is not set in the environment variables. Please set it and try again.")
    st.stop()

#LLM initialization
llm = HuggingFaceEndpoint(
    repo_id="mistralai/Mixtral-8x7B-Instruct-v0.1",
    huggingfacehub_api_token=HF_API_KEY, 
    temperature=0.7,
    max_new_tokens=500
)

prompt_template=""" 
You are data quality expert. Given the following data quality issues for a csv file :
{issues}
Please provide a detailed analysis of the issues
1.what is the issue and why it matters.
2.Suggest fixes for each issue (e.g drop rows, impute values,remove duplicates, etc.)
3.how it will impact by applying the fix
Return response in a clear conscise language suitable for a data scientist
"""
# Define the prompt templates
prompt= PromptTemplate(template=prompt_template,input_variables=["issues"])
llm_chain = LLMChain(llm=llm, prompt=prompt)


# function to do anlysis on the uploaded file

def inspect_data(df):

    issues = {}

    #missing values
    missing = df.isnull().sum()
    missing= missing[missing > 0]
    if not missing.empty:
        issues['Missing Values'] = missing.to_dict()


    #duplicate rows
    duplicates = df.duplicated().sum() 
    print(f"Number of duplicate rows: {duplicates}")
    if duplicates > 0:
        issues['Duplicate Rows'] = int(duplicates)


    # outliers we can use IOR methods to underrstand outliers
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    outliers = {}
    for col in numeric_cols:         #each column we calculate the IQR and filter outliers so that we can get the count of outliers
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        outlier_count = df[(df[col] <lower_bound) | (df[col] >upper_bound)].shape[0] # less than 25th percentile or greater than 75th percentile
        if outlier_count > 0:
            outliers[col] = outlier_count
    if outliers:
        issues['Outliers'] = outliers


    # mismatched data types
    dtypes = df.dtypes.to_dict()
    schema_issues = {}
    for col,dtype in dtypes.items():
        is_string=True
        is_numeric= True

        for value in df[col]:
            if pd.notnull(value):
                if not isinstance(value, (int, float)):
                    is_numeric = False
                    break
                if not isinstance(value, str):
                    is_string = False
                    break
        if is_string and dtype != 'object':
            schema_issues[col] = f"Expected string type but found {dtype}"
        elif is_numeric and dtype != 'float64' and dtype != 'int64':
            schema_issues[col] = f"Expected numeric type but found {dtype}" 
    if schema_issues:
        issues['Schema Issues'] = schema_issues
    



    return issues


# cleaning function 

def clean_data(df,fix_duplicates,fix_outliers,fix_missing):
    cleaned_df=df.copy()
    logs=[]

    #missing values for both numeric and category 
    if fix_missing=="Drop":
        initial_rows=len(cleaned_df)
        cleaned_df=cleaned_df.dropna()
        logs.append(f"Dropped{initial_rows-len(cleaned_df)} rows with missing values")
    elif fix_missing=="Impute":
        for col in cleaned_df.columns:
            if cleaned_df[col].isnull().any:
                if cleaned_df[col].dtype in ["int64","float64"]:
                    mean_val=cleaned_df[col].mean()
                    cleaned_df[col]=cleaned_df[col].fillna(mean_val)
                    logs.append(f"imputed missing values in {col} with mean {mean_val}")
                else:
                    mode_val=cleaned_df[col].mode()[0]
                    cleaned_df[col]=cleaned_df[col].fillna(mode_val)
                    logs.append(f"imputed missing values in {col} with mode {mode_val}")
    st.success("Handled misisng values ")


    #Handling ouliers
    if fix_outliers:
        for col in cleaned_df.select_dtypes(include=[np.number]).columns:
            Q1=cleaned_df[col].quantile(0.25)
            Q3=cleaned_df[col].quantile(0.75)
            IQR=Q3-Q1

            lower_bound=Q1-1.5*IQR
            higher_bound=Q3+1.5*IQR

            initial_rows=len(cleaned_df)
            cleaned_df=cleaned_df[(cleaned_df[col] >=lower_bound) & (cleaned_df[col] <=higher_bound)]
            logs.append(f"removed {initial_rows-len(cleaned_df)} rows with outliers in {col}")
    st.success("handles outliers")       






    #handle dupes
    if fix_duplicates:
        initial_rows=len(cleaned_df)
        cleaned_df=cleaned_df.drop_duplicates()
        logs.append(f"removed {initial_rows-len(cleaned_df)} duplicated rows ")

        st.success("duplicates got removed")
        


    return cleaned_df,logs


# function create docx

def generate_report(issues,llm_response,logs) :

    doc=Document()
    doc.add_heading("Data Quality Report",level=0)


    doc.add_heading("issues",level=1)
    for issue_type,detailes in issues.items():
        doc.add_heading(issue_type,level=2)
        doc.add_paragraph(str(detailes))

    doc.add_heading("Ai Anlysis",level=1)
    doc.add_paragraph(llm_response)

    doc.add_heading("Logs",level=1)
    for log in logs:
        doc.add_paragraph(log)
    
    output=io.BytesIO()
    doc.save(output)
    return output.getvalue()




#streamlit app interface

st.title("Ai Agent for prelimanary Data analysis")

uploaded_file = st.file_uploader("Upload a DOCX file", type="csv")

if uploaded_file is not None:

    df=pd.read_csv(uploaded_file)
    st.write("Dataframe loaded successfully!",st.write(df.dtypes))
    st.dataframe(df)

    # Inspect the data
    issues = inspect_data(df)

    
    

    #Paring the data to LLM
    if issues:
        st.write("Issues found in the data:")
        st.json(issues)

        with st.spinner("Analyzing data..."):
            
            llm_response = llm_chain.run(issues=issues)
            st.write(llm_response)


    # fixing the problems
    st.write("##  Apply Fixes ")
    fix_missing=st.selectbox("handle miss values",["None","Drop","Impute"])
    fix_duplicates=st.checkbox("Remove duplicates")
    fix_outliers=st.checkbox("Remove outliers")



    if st.button("Apply"):

        cleaned_df,logs=clean_data(df,fix_duplicates,fix_outliers,fix_missing)

        for log in logs:
            st.write(log)
        st.write("cleaned data preview")
        st.dataframe(cleaned_df.head(10))


    

        report_data =generate_report(issues,llm_response,logs)
        st.download_button(
            label="Download_report(DOCX)",
            data=report_data,
            file_name="data_quality_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
            
else:
    st.write("No data quality issues")






    




      


